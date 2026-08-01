#!/usr/bin/env python3
"""Generate a synthetic VNI corpus from the public-domain Vietnamese already collected.

Real VNI documents are scarce: 28 Vietnamese government domains produced two, and the
diaspora sources that have more are copyrighted. So the source text here is the corpus's
own TCVN3 transcripts — real Vietnamese legal prose, already public under Article 15,
re-encoded into VNI.

.. warning::
   **This measures self-consistency, not correctness.** The text is encoded with
   ``scripts/vni.py`` and any parser is scored against a table derived from the same
   two documents that table came from. It will find *missing* entries in another
   implementation — which is the point, viparse's VNI table has 6 where roughly fifty
   are needed — but it cannot show that either table is right.

   No headline number comes from this set. ``METRIC.md`` requires the synthetic and
   public-domain subsets to be reported separately, and this is why.

Documents are written as ``.docx`` with the run font set to ``VNI-Times``, so a parser
sees the same font signal a real VNI document carries and the extraction and detection
layers are exercised, not just the conversion table.

    python3 scripts/make_synthetic.py --count 12
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from vni import Unencodable, convert, encode  # noqa: E402

ROOT = Path(__file__).resolve().parent.parent
TRUTH = ROOT / "ground-truth"
OUT = ROOT / "corpus" / "synthetic"
SYNTH_TRUTH = ROOT / "ground-truth-synthetic"


def source_rows() -> dict[str, list[str]]:
    """Every row under `## Files`, keyed by stem — the chain a generated file inherits.

    A synthetic document is only redistributable because its source text is. Copying the
    publisher and legal basis across from the source row makes that chain explicit rather
    than leaving a generated file with a provenance entry that says only "generated".
    """
    rows: dict[str, list[str]] = {}
    lines = (ROOT / "PROVENANCE.md").read_text(encoding="utf-8").split("\n")
    start = next(i for i, ln in enumerate(lines) if ln.strip().lower() == "## files")
    for line in lines[start + 1 :]:
        if line.strip().startswith("#"):
            break
        if not line.strip().startswith("|"):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        name = cells[0].strip("`")
        if name in ("file",) or set(name) <= set("-: "):
            continue
        rows[name.removesuffix(".doc").removesuffix(".docx")] = cells
    return rows


def write_provenance(entries: list[tuple[str, list[str]]]) -> None:
    """Replace the synthetic rows in PROVENANCE.md's `## Files` table.

    Written from the entries actually produced, in the same pass that produced them. The
    sibling `fetch_corpus.py` once recorded provenance only after its whole run finished,
    and an interruption left 59 files with no entry at all — the exact state the
    validator exists to catch.
    """
    path = ROOT / "PROVENANCE.md"
    lines = path.read_text(encoding="utf-8").split("\n")
    kept = [ln for ln in lines if "`synthetic-vni-" not in ln]
    start = next(i for i, ln in enumerate(kept) if ln.strip().lower() == "## files")
    last = max(
        i for i, ln in enumerate(kept[start:], start) if ln.strip().startswith("|")
    )
    new = []
    for stem, src in entries:
        source, retrieved, publisher, basis = src[1], src[2], src[3], src[4]
        new.append(
            f"| `{stem}.docx` | generated from `{src[0].strip('`')}` ({source}) | "
            f"{retrieved} | {publisher} | {basis} | vni | ready | "
            f"synthetic: text of the source transcript re-encoded to VNI by "
            f"scripts/make_synthetic.py, font VNI-Times; round trip verified per line |"
        )
    path.write_text("\n".join(kept[: last + 1] + new + kept[last + 1 :]), encoding="utf-8")


def ready_stems() -> set[str]:
    """Stems whose transcript PROVENANCE.md marks `ready`.

    Read from the provenance table rather than from the directory listing. A transcript
    can be demoted — `2004-1ED4D_Francois_Godement` was, for mixing French into
    Vietnamese — and a generator that globs the directory silently re-admits it, which
    is how a known-bad transcript ends up supplying a fifth of a synthetic subset.
    """
    stems: set[str] = set()
    for line in (ROOT / "PROVENANCE.md").read_text(encoding="utf-8").split("\n"):
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if len(cells) < 7 or cells[6] != "ready":
            continue
        stems.add(cells[0].strip("`").removesuffix(".doc").removesuffix(".docx"))
    return stems


def write_docx(path: Path, lines: list[str]) -> None:
    import docx
    from docx.shared import Pt

    document = docx.Document()
    for line in lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        # The font is the signal a parser is meant to detect; without it this would
        # test the conversion table in isolation and nothing upstream of it.
        run.font.name = "VNI-Times"
        run.font.size = Pt(13)
    document.save(str(path))


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--count", type=int, default=12, help="documents to generate")
    ap.add_argument("--max-lines", type=int, default=120, help="lines taken per source")
    args = ap.parse_args()

    ready = ready_stems()
    sources = [t for t in sorted(TRUTH.glob("*.txt")) if t.stem in ready][: args.count]
    if not sources:
        print("no `ready` transcripts to draw from", file=sys.stderr)
        return 1
    skipped = len(list(TRUTH.glob("*.txt"))) - len([t for t in TRUTH.glob("*.txt") if t.stem in ready])
    if skipped:
        print(f"  {skipped} transcript(s) not `ready`, not used as a source", file=sys.stderr)

    OUT.mkdir(parents=True, exist_ok=True)
    SYNTH_TRUTH.mkdir(parents=True, exist_ok=True)

    made = mismatches = 0
    dropped: dict[str, int] = {}
    rows = source_rows()
    entries: list[tuple[str, list[str]]] = []
    for source in sources:
        text = source.read_text(encoding="utf-8")
        lines = [ln for ln in text.split("\n") if ln.strip()][: args.max_lines]
        if not lines:
            continue

        # Drop lines the table cannot express, rather than the whole document. ẳ and ẵ
        # have no observed VNI sequence, so a line containing either is not something
        # this generator can honestly produce; the rest of the document still is.
        kept: list[str] = []
        encoded: list[str] = []
        for line in lines:
            try:
                surface = encode(line)
            except Unencodable as gap:
                dropped[str(gap)] = dropped.get(str(gap), 0) + 1
                continue
            # Verify the round trip line by line. A generator that does not reproduce its
            # own input is measuring itself, and the failure would look like a parser bug.
            if convert(surface) != line:
                mismatches += 1
                continue
            kept.append(line)
            encoded.append(surface)

        if not encoded:
            continue
        lines = kept

        stem = f"synthetic-vni-{source.stem}"
        entries.append((stem, rows[source.stem]))
        write_docx(OUT / f"{stem}.docx", encoded)
        (SYNTH_TRUTH / f"{stem}.txt").write_text("\n".join(lines) + "\n", encoding="utf-8")
        made += 1

    write_provenance(entries)
    print(f"  {made} document(s) written to {OUT.relative_to(ROOT)}", file=sys.stderr)
    print(f"  {len(entries)} provenance row(s) written to PROVENANCE.md", file=sys.stderr)
    if mismatches:
        print(f"  {mismatches} line(s) skipped on a failed round trip", file=sys.stderr)
    if dropped:
        gaps = ", ".join(f"{c}×{n}" for c, n in sorted(dropped.items(), key=lambda kv: -kv[1]))
        print(f"  line(s) dropped, no observed VNI sequence for: {gaps}", file=sys.stderr)
    print("  Synthetic — never a headline number. See METRIC.md.", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
